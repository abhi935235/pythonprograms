from tkinter import *

#from docutils.nodes import entry
from docx import Document
import matplotlib.pyplot as plt

root = Tk()
root.title("GUI")
root.geometry("2000x600+600+500")
root.resizable(height=False, width=False)

total = []
total1 = []
sample = 1
e1 = []
e2 = []
e3 = []
e4 = []
e5 = []
e6 = []
e7 = []
e8 = []
e9= []
e10 = []
e11 = []
n=3

def avg():
    tot_sum = 0
    for i in total:
        tot_sum += i
    avg = tot_sum/len(total)
    print(avg)
    text = Text(root)
    text.insert(END,avg)
    text.grid(row=2, column=7)

# def new():
#     global ent1
#     global ent2
#     lb1 = Label(root,text="Length")
#     ent1 = Entry(root)
#     lb1.pack()
#     ent1.pack()
#     btn = Button(root,text="add",command=show)
#     btn.pack()

def new1():
    global ent2
    global ent3
    global ent4
    #global n
    #n+=1
   # for i in range(n):

n=3

def xy():
    c1 = 0
    c2 = 0
    c3 = 0
    c4 = 0
    c5 = 0
    c6 = 0
    c7 = 0
    for i in range(e1):
        c1 += i
    for i in range(e2):
        c2 += i
    for i in range(e3):
        c3 += i
    for i in range(e4):
        c4 += i
    for i in range(e5):
        c5 += i
    for i in range(e6):
        c6 += i
    for i in range(e7):
        c7 += i
    mean = c1+c2+c3+c4+c5+c6+c7/7
    print(mean)
    text = Text(root)
    text.insert(END, mean)
    text.grid(row=2, column=5)

def add():
    global n
    global en2
    global en3
    global en4
    global en5
    global en6
    global en7
    global en8
    global en9
    global en10
    global en11
    global en12
    global en13
    n+=1
    for i in range(n):

        en2 = Entry(root)
        en2.grid(row=n, column=1)
        en3 = Entry(root)
        en3.grid(row=n, column=2)
        en4 = Entry(root)
        en4.grid(row=n, column=3)
        en5 = Entry(root)
        en5.grid(row=n, column=4)
        en6 = Entry(root)
        en6.grid(row=n, column=5)
        en7 = Entry(root)
        en7.grid(row=n, column=6)
        en8 = Entry(root)
        en8.grid(row=n, column=7)
        en9 = Entry(root)
        en9.grid(row=n, column=8)
        en10 = Entry(root)
        en10.grid(row=n, column=9)
        en11 = Entry(root)
        en11.grid(row=n, column=10)
        en12 = Entry(root)
        en12.grid(row=n, column=11)
        en13 = Entry(root)
        en13.grid(row=n, column=12)
        btn = Button(root, text="add", command=show)
        btn.grid(row=n, column=13)







def show():
    global sample
    en11.insert(0,str(float(en2.get())+float(en3.get())+float(en4.get())+float(en5.get())+float(en6.get())+float(en7.get())+float(en8.get())+float(en9.get())+float(en10.get())/9))
    en13.insert(0, str(float(en11.get())+int(en12.get())))
    row = table.add_row()
    # total.append(int(ent1.get()))
    e1.append(float(en2.get()))
    e2.append(float(en3.get()))
    e3.append(float(en4.get()))
    e4.append(float(en5.get()))
    e5.append(float(en6.get()))
    e6.append(float(en7.get()))
    e7.append(float(en8.get()))
    e8.append(float(en9.get()))
    e9.append(float(en10.get()))
    e10.append(float(en11.get()))
    e11.append(float(en12.get()))
    cell = table.rows[len(e1)].cells
    cell[0].text = str(sample)
    # cell[1].text = ent1.get()
    cell[1].text = en2.get()
    cell[2].text = en3.get()
    cell[3].text = en4.get()
    cell[4].text = en5.get()
    cell[5].text = en6.get()
    cell[6].text = en7.get()
    cell[7].text = en8.get()
    cell[8].text = en9.get()
    cell[9].text = en10.get()
    #cell[3].text = str(int(ent2.get())/int(ent3.get()))
    print(e1)
    print(e2)
    print(e3)
    print(e4)
    print(e5)
    print(e6)
    print(e7)
    print(e8)
    print(e9)
    sample+=1
    # new()
    add()

def quit():
    doc.save("mod3.docx")
    root.destroy()

def graph():
    global total1
    global total
    doc.save("mod3.docx")
    ndoc = Document('mod3.docx')
    table = ndoc.tables[0]

    sample_n = []
    data = []
    data1 = []
    data2 = []
    data3 = []
    data4 = []
    data5 = []
    data6 = []
    data7 = []
    data8 = []

    i = 0
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)
            if i==0:
                i+=1
                continue
            elif i%4 == 0:
                sample_n.append(int(cell.text))
                i+=1
            elif i%7 == 0:
                data.append(int(cell.text))
                i+=1

    print(len(sample_n))
    print(len(data))
    print(len(data1))
    print(len(data2))
    print(len(data3))
    print(len(data4))
    print(len(data5))
    print(len(data6))
    print(len(data7))
    print(len(data8))

    graphvar = []
    if len(sample_n) == 0:
        for i in range(len(e1)):
            graphvar.append((i + 1, e1[i]))
        for i in range(len(e1)):
            graphvar.append((i + 1, e2[i]))
        for i in range(len(e1)):
            graphvar.append((i + 1, e3[i]))
        for i in range(len(e1)):
            graphvar.append((i + 1, e4[i]))
        for i in range(len(e1)):
            graphvar.append((i + 1, e5[i]))
        for i in range(len(e1)):
            graphvar.append((i + 1, e6[i]))
        for i in range(len(e1)):
            graphvar.append((i + 1, e7[i]))
        for i in range(len(e1)):
            graphvar.append((i + 1, e8[i]))
        ind = []
        value = []
        for index,val in graphvar:
            ind.append(index)
            value.append(val)
        plt.plot(ind,value)
        plt.xlabel(['sample'])
        plt.ylabel(['cuts'])
        plt.savefig('C:/Users/veran/Desktop/final/graph3.png')
        doc.add_picture('C:/Users/veran/Desktop/final/graph3.png')
        plt.show()
    else:
        plt.plot(sample_n,data)
        plt.xlabel(['sample'])
        plt.ylabel(['cuts'])
        plt.show()
    print(graphvar)

doc = Document()
doc.add_heading("Measurement data of input wire electrical resistance",0)
table = doc.add_table(rows=1,cols=11)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Sample number"
hdr_cells[1].text = "cu1"
hdr_cells[2].text = "cut2"
hdr_cells[3].text = "cut3"
hdr_cells[4].text = "cut4"
hdr_cells[5].text = "cut5"
hdr_cells[6].text = "cut6"
hdr_cells[7].text = "cut7"
hdr_cells[8].text = "mean"
hdr_cells[9].text = "cna"
hdr_cells[10].text = "total"



btn1 = Button(root,text="add Data",command=add)
btn1.grid(row=1,column=1)
# btn6 = Button(root,text="add Resistance",command=new1)
# btn6.pack()
btn4 = Button(root,text="show graph",command=graph)
btn4.grid(row=1,column=6)
btn3 = Button(root,text="quit",command=quit)
btn3.grid(row=1, column=7)
l2 = Label(root,text="cu1")
l2.grid(row=2,column=1)
l3 = Label(root, text="cut2")
l3.grid(row=2, column=2)
l4 = Label(root, text="cut3")
l4.grid(row=2, column=3)
l5 = Label(root, text="cut4")
l5.grid(row=2, column=4)
l6 = Label(root, text="cut5")
l6.grid(row=2, column=5)
l7 = Label(root, text="cut6")
l7.grid(row=2, column=6)
l8 = Label(root, text="cut7")
l8.grid(row=2, column=7)
l9 = Label(root, text="cut8")
l9.grid(row=2, column=8)
l10 = Label(root, text="cut9")
l10.grid(row=2, column=9)
l11 = Label(root, text="mean")
l11.grid(row=2, column=10)
l12 = Label(root, text="Containers & Accessories")
l12.grid(row =2, column=11)
l13 = Label(root, text ="Addition")
l13.grid(row = 2, column=12)
root.mainloop()
